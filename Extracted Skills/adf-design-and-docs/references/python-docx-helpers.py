"""
ADF Haiti Word Document Helpers (python-docx)

Ready-to-use functions that apply ADF's brand standards to any python-docx document.
Copy the functions you need into your generation script, or import this module directly.

Tables use Word's built-in `GridTable4` style (neutral grayscale). ADF does NOT apply
per-cell shading, borders, or margins, and does NOT force bold or color on runs inside
table cells. The style does that work.

Updated April 2026 (iteration 2): corrected from user-revised reference documents.
- Headings use "Aptos SemiBold" (not Aptos Bold)
- H1: 16pt SemiBold, spacing 40/0 (was 200/80)
- H2: 12pt SemiBold, spacing 40/0 (was 160/80)
- H3: 11pt SemiBold, spacing 40/0 (was 120/40)
- Body line spacing: 276 twips (1.15x), not single
- Horizontal rules: REMOVED (forbidden by ADF)
- Table row height minimum: 0.22" (317 DXA)
- Footer distance: 432 DXA (0.3")

Updated May 2026 (adf-design-and-docs merge): three brand changes.
- Page margins: 0.75" all sides on EVERY section (was 1.0"). Applied to
  generated docs and to documents opened from a .dotx template.
- Content width: 10080 DXA (8.5" - 2x0.75"). Column presets rescaled to 10080.
- Footer: page number only ("Paj N" / "Page N"). The "Konfidansyel/Confidential"
  label has been removed entirely.
- Haitian Creole / French text must carry correct diacritics (e, o, a with
  grave/acute accents). See references/haitian-creole-orthography.md.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from copy import deepcopy
from lxml import etree


# ============================================================================
# ADF Brand Constants
# ============================================================================

ADF_COLORS = {
    "raisin_black":   "231F20",
    "davy_grey":      "58585B",
    "prussian_blue":  "0A3D50",
    "smooth_green":   "3EAC7A",
    "light_sea_green":"41AAA3",
    "neptune_cyan":   "6EB7B2",
    "pale_blue_lily": "D8E7E6",
    "white":          "FFFFFF",
    "amber":          "FFC107",
    "orange":         "FF9800",
    "red":            "F44336",
    "dark_red":       "B71C1C",
}

ADF_FONT = "Aptos"
ADF_FONT_SEMIBOLD = "Aptos SemiBold"  # for headings (lighter than Bold)
ADF_FONT_FALLBACK = "Inter"

# US Letter, 0.75" margins on all sides (ADF default, May 2026)
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
DEFAULT_MARGIN_DXA = 1080   # 0.75" (was 1440 / 1.0")
HEADER_DISTANCE_DXA = 720
FOOTER_DISTANCE_DXA = 432   # 0.3"
CONTENT_WIDTH_DEFAULT = 10080  # 8.5" page - 2 x 0.75" margins
CONTENT_WIDTH_NARROW = 10080   # retained alias; equals default

# Minimum table row height: 0.22" = 316.8 DXA, rounded up to 317
MIN_ROW_HEIGHT_DXA = 317

# Page-number label only. The confidentiality label was removed (May 2026).
# Diacritics preserved where the language requires them (es: Pagina -> Pagina).
PAGE_LABELS = {
    "ht": "Paj ",
    "en": "Page ",
    "fr": "Page ",
    "es": "Página ",
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ============================================================================
# Built-in Table Style Definitions (injected into styles.xml)
# ============================================================================

TABLE_NORMAL_XML = f'''<w:style xmlns:w="{W_NS}" w:type="table" w:default="1" w:styleId="TableNormal">
  <w:name w:val="Normal Table"/>
  <w:uiPriority w:val="99"/>
  <w:semiHidden/>
  <w:unhideWhenUsed/>
  <w:tblPr>
    <w:tblInd w:w="0" w:type="dxa"/>
    <w:tblCellMar>
      <w:top w:w="0" w:type="dxa"/>
      <w:left w:w="108" w:type="dxa"/>
      <w:bottom w:w="0" w:type="dxa"/>
      <w:right w:w="108" w:type="dxa"/>
    </w:tblCellMar>
  </w:tblPr>
</w:style>'''

GRID_TABLE_4_XML = f'''<w:style xmlns:w="{W_NS}" w:type="table" w:styleId="GridTable4">
  <w:name w:val="Grid Table 4"/>
  <w:basedOn w:val="TableNormal"/>
  <w:uiPriority w:val="49"/>
  <w:pPr>
    <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
  </w:pPr>
  <w:tblPr>
    <w:tblStyleRowBandSize w:val="1"/>
    <w:tblStyleColBandSize w:val="1"/>
    <w:tblBorders>
      <w:top w:val="single" w:sz="4" w:space="0" w:color="868686"/>
      <w:left w:val="single" w:sz="4" w:space="0" w:color="868686"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="868686"/>
      <w:right w:val="single" w:sz="4" w:space="0" w:color="868686"/>
      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="868686"/>
      <w:insideV w:val="single" w:sz="4" w:space="0" w:color="868686"/>
    </w:tblBorders>
  </w:tblPr>
  <w:tblStylePr w:type="firstRow">
    <w:rPr>
      <w:b/>
      <w:bCs/>
      <w:color w:val="FFFFFF"/>
    </w:rPr>
    <w:tblPr/>
    <w:tcPr>
      <w:tcBorders>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="363636"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="363636"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="363636"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="363636"/>
        <w:insideH w:val="nil"/>
        <w:insideV w:val="nil"/>
      </w:tcBorders>
      <w:shd w:val="clear" w:color="auto" w:fill="363636"/>
    </w:tcPr>
  </w:tblStylePr>
  <w:tblStylePr w:type="lastRow">
    <w:rPr><w:b/><w:bCs/></w:rPr>
    <w:tblPr/>
    <w:tcPr>
      <w:tcBorders>
        <w:top w:val="double" w:sz="4" w:space="0" w:color="363636"/>
      </w:tcBorders>
    </w:tcPr>
  </w:tblStylePr>
  <w:tblStylePr w:type="firstCol"><w:rPr><w:b/><w:bCs/></w:rPr></w:tblStylePr>
  <w:tblStylePr w:type="lastCol"><w:rPr><w:b/><w:bCs/></w:rPr></w:tblStylePr>
  <w:tblStylePr w:type="band1Vert">
    <w:tblPr/>
    <w:tcPr><w:shd w:val="clear" w:color="auto" w:fill="D6D6D6"/></w:tcPr>
  </w:tblStylePr>
  <w:tblStylePr w:type="band1Horz">
    <w:tblPr/>
    <w:tcPr><w:shd w:val="clear" w:color="auto" w:fill="D6D6D6"/></w:tcPr>
  </w:tblStylePr>
</w:style>'''


# ============================================================================
# Core Style Application
# ============================================================================

def apply_adf_styles(doc):
    """
    Apply ADF brand styles to a new or existing Document.

    Sets document-wide defaults (Aptos body, Raisin Black text), overrides
    Heading 1/2/3/Title to ADF specs, and injects the `TableNormal` and
    `GridTable4` table styles into styles.xml if missing.
    """
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = ADF_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(ADF_COLORS["raisin_black"])

    # Force Aptos into rFonts and strip theme font refs (theme refs cause Calibri fallback)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ADF_FONT)
    rfonts.set(qn("w:hAnsi"), ADF_FONT)
    rfonts.set(qn("w:cs"), ADF_FONT)
    rfonts.set(qn("w:eastAsia"), ADF_FONT)
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        key = qn(theme_attr)
        if key in rfonts.attrib:
            del rfonts.attrib[key]

    # Normal paragraph spacing: after 80 DXA, line spacing 276 (1.15x)
    _set_style_spacing(normal, space_before=0, space_after=80, line=276, line_rule="auto")

    # Heading 1: Aptos SemiBold 16pt Prussian Blue, spacing 40/0, single line
    _configure_heading(styles["Heading 1"], size_pt=16, color_hex=ADF_COLORS["prussian_blue"],
                       space_before=40, space_after=0, line=240, line_rule="auto")

    # Heading 2: Aptos SemiBold 12pt Prussian Blue, spacing 40/0, single line
    _configure_heading(styles["Heading 2"], size_pt=12, color_hex=ADF_COLORS["prussian_blue"],
                       space_before=40, space_after=0, line=240, line_rule="auto")

    # Heading 3: Aptos SemiBold 11pt Davy Grey, spacing 40/0, single line
    _configure_heading(styles["Heading 3"], size_pt=11, color_hex=ADF_COLORS["davy_grey"],
                       space_before=40, space_after=0, line=240, line_rule="auto")

    # Title: Aptos SemiBold 28pt Prussian Blue, spacing 240/0
    if "Title" in styles:
        _configure_heading(styles["Title"], size_pt=28, color_hex=ADF_COLORS["prussian_blue"],
                           space_before=240, space_after=0)

    # Subtitle: Aptos 14pt Regular Davy Grey, spacing 0/200
    if "Subtitle" in styles:
        subtitle = styles["Subtitle"]
        subtitle.font.name = ADF_FONT
        subtitle.font.size = Pt(14)
        subtitle.font.color.rgb = RGBColor.from_string(ADF_COLORS["davy_grey"])
        subtitle.font.bold = False
        _set_style_spacing(subtitle, space_before=0, space_after=200)

    # Strip theme font refs from docDefaults (prevents Calibri fallback globally)
    _strip_doc_defaults_theme_fonts(doc)

    # Inject TableNormal + GridTable4 into styles.xml
    _inject_table_styles(doc)

    # Page layout: 0.75" margins on EVERY section. Looping over all sections
    # (not just sections[0]) forces the ADF geometry even on documents opened
    # from a .dotx template or containing multiple sections.
    force_adf_margins(doc)

    return doc


def force_adf_margins(doc, margin_dxa=DEFAULT_MARGIN_DXA):
    """Force ADF page geometry (0.75" margins) on all sections of a document.

    Call this on any document, including one opened from a .dotx template, to
    override the template's margins. Sets US Letter page size, uniform
    0.75" margins, and ADF header/footer distances on every section.
    """
    for section in doc.sections:
        section.page_width = Twips(PAGE_WIDTH_DXA)
        section.page_height = Twips(PAGE_HEIGHT_DXA)
        section.top_margin = Twips(margin_dxa)
        section.bottom_margin = Twips(margin_dxa)
        section.left_margin = Twips(margin_dxa)
        section.right_margin = Twips(margin_dxa)
        section.header_distance = Twips(HEADER_DISTANCE_DXA)
        section.footer_distance = Twips(FOOTER_DISTANCE_DXA)
    return doc


def _strip_doc_defaults_theme_fonts(doc):
    """Remove theme font references from docDefaults so Aptos is used globally.

    python-docx documents inherit Calibri from the default theme's minorFont.
    Stripping w:asciiTheme etc. from rPrDefault forces Word to use the explicit
    rFonts values (Aptos) we set on the Normal style instead.
    """
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is not None:
        rpr_default = doc_defaults.find(qn("w:rPrDefault"))
        if rpr_default is not None:
            rpr = rpr_default.find(qn("w:rPr"))
            if rpr is not None:
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme",
                                       "w:cstheme", "w:eastAsiaTheme"):
                        key = qn(theme_attr)
                        if key in rfonts.attrib:
                            del rfonts.attrib[key]
                    # Set explicit Aptos in docDefaults too
                    rfonts.set(qn("w:ascii"), ADF_FONT)
                    rfonts.set(qn("w:hAnsi"), ADF_FONT)
                    rfonts.set(qn("w:cs"), ADF_FONT)
                    rfonts.set(qn("w:eastAsia"), ADF_FONT)


def _inject_table_styles(doc):
    """Ensure TableNormal and GridTable4 exist in styles.xml."""
    styles_element = doc.styles.element
    existing_ids = set()
    for style_el in styles_element.findall(qn("w:style")):
        sid = style_el.get(qn("w:styleId"))
        if sid:
            existing_ids.add(sid)

    if "TableNormal" not in existing_ids:
        styles_element.append(etree.fromstring(TABLE_NORMAL_XML))
    if "GridTable4" not in existing_ids:
        styles_element.append(etree.fromstring(GRID_TABLE_4_XML))


def _configure_heading(style, size_pt, color_hex, space_before, space_after,
                       line=None, line_rule=None):
    """Set font, size, color, spacing on a heading style. Uses Aptos SemiBold."""
    style.font.name = ADF_FONT_SEMIBOLD
    style.font.size = Pt(size_pt)
    style.font.bold = False  # SemiBold is a separate font face, not bold flag
    style.font.color.rgb = RGBColor.from_string(color_hex)

    # Force Aptos SemiBold into rFonts and strip theme font refs
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), ADF_FONT_SEMIBOLD)
    # Remove theme font attributes (these override explicit names and cause Calibri fallback)
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme"):
        key = qn(theme_attr)
        if key in rfonts.attrib:
            del rfonts.attrib[key]

    # Remove any explicit bold elements (SemiBold weight comes from font face)
    for tag in ("w:b", "w:bCs"):
        el = rpr.find(qn(tag))
        if el is not None:
            rpr.remove(el)

    _set_style_spacing(style, space_before, space_after, line=line, line_rule=line_rule)

    # Keep headings with the content that follows them (prevents orphaned headings)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.keep_together = True


def _set_style_spacing(style, space_before, space_after, line=None, line_rule=None):
    """Set paragraph spacing on a style, including optional line spacing."""
    ppr = style.element.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), str(space_before))
    spacing.set(qn("w:after"), str(space_after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
    if line_rule is not None:
        spacing.set(qn("w:lineRule"), line_rule)


# ============================================================================
# Header and Footer
# ============================================================================

def add_adf_header(doc, document_type_label):
    """
    Add standard ADF header: ADF Haiti | [Document Type]
    Prussian Blue bottom border.
    """
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.text = ""

    r1 = p.add_run("ADF Haiti")
    r1.font.name = ADF_FONT
    r1.font.size = Pt(9)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor.from_string(ADF_COLORS["prussian_blue"])

    r2 = p.add_run("  |  ")
    r2.font.name = ADF_FONT
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(ADF_COLORS["davy_grey"])

    r3 = p.add_run(document_type_label)
    r3.font.name = ADF_FONT
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor.from_string(ADF_COLORS["davy_grey"])

    # No border line in header (clean design per ADF letterhead reference)


def add_adf_footer(doc, language="ht"):
    """
    Add standard ADF footer: page number only, e.g. "Paj 1" (ht) / "Page 1" (en).
    Centered, Aptos 8pt Davy Grey. The confidentiality label was removed (May 2026).
    No border line (clean design per ADF letterhead reference).
    """
    page_label = PAGE_LABELS.get(language, PAGE_LABELS["ht"])

    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""

    r1 = p.add_run(page_label)
    r1.font.name = ADF_FONT
    r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor.from_string(ADF_COLORS["davy_grey"])

    _add_page_number_field(p)


def _add_page_number_field(paragraph):
    """Insert a PAGE field into the paragraph."""
    fld_begin = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), ADF_FONT)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    rpr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ADF_COLORS["davy_grey"])
    rpr.append(color)
    fld_begin.append(rpr)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    fld_begin.append(fld_char)
    paragraph._p.append(fld_begin)

    instr = OxmlElement("w:r")
    instr.append(deepcopy(rpr))
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    instr.append(instr_text)
    paragraph._p.append(instr)

    fld_end = OxmlElement("w:r")
    fld_end.append(deepcopy(rpr))
    fld_end_char = OxmlElement("w:fldChar")
    fld_end_char.set(qn("w:fldCharType"), "end")
    fld_end.append(fld_end_char)
    paragraph._p.append(fld_end)


def _add_paragraph_border(paragraph, position, color_hex, size, space, style="single"):
    """Add a border (top/bottom/left/right) to a paragraph."""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), style)
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color_hex)
    pbdr.append(border)


# ============================================================================
# Tables (GridTable4)
# ============================================================================

def build_adf_table(doc, headers, rows, column_widths=None, total_width=None,
                    repeat_header=True, last_row_total=False,
                    first_column_emphasis=True):
    """
    Build a data table styled via Word's built-in GridTable4 style.

    The table references GridTable4 and applies no per-cell shading, borders,
    or margins. Header-row formatting, zebra banding, and borders are rendered
    by the style definition injected by apply_adf_styles().

    Args:
        doc: python-docx Document (must have had apply_adf_styles() called)
        headers: list[str] (first row)
        rows: list[list[str]] (data rows)
        column_widths: list[int] in DXA. If None, split evenly.
        total_width: int in DXA. Defaults to 9360.
        repeat_header: repeat header row on page breaks
        last_row_total: enable GridTable4 lastRow styling (totals row)
        first_column_emphasis: bold left column (label convention)

    Returns the python-docx Table object.
    """
    if total_width is None:
        total_width = CONTENT_WIDTH_DEFAULT
    if column_widths is None:
        col_count = len(headers)
        even = total_width // col_count
        column_widths = [even] * (col_count - 1) + [total_width - even * (col_count - 1)]

    assert sum(column_widths) == total_width, (
        f"Column widths {column_widths} (sum={sum(column_widths)}) != total_width={total_width}"
    )
    assert len(column_widths) == len(headers)
    for row in rows:
        assert len(row) == len(headers), f"Row length {len(row)} != header count {len(headers)}"

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    table.allow_autofit = False

    _set_gridtable4_tblPr(table, repeat_header=repeat_header,
                          last_row_total=last_row_total,
                          first_column_emphasis=first_column_emphasis)
    _set_tbl_grid(table, column_widths)

    # Header row : only set sz and white color; bold + font name come from
    # GridTable4 firstRow conditional formatting and Normal style inheritance.
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_width(cell, column_widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _strip_cell_direct_formatting(cell)
        p = cell.paragraphs[0]
        _clear_paragraph_runs(p)
        p.paragraph_format.space_before = Twips(0)
        p.paragraph_format.space_after = Twips(0)
        run = p.add_run(header_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(ADF_COLORS["white"])

    if repeat_header:
        tr = table.rows[0]._tr
        trpr = tr.find(qn("w:trPr"))
        if trpr is None:
            trpr = OxmlElement("w:trPr")
            tr.insert(0, trpr)
        tbl_header = OxmlElement("w:tblHeader")
        trpr.append(tbl_header)

    # Data rows : only set sz (10pt, differs from body 11pt); font name and
    # color inherit from Normal style.
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[1 + row_idx].cells[col_idx]
            _set_cell_width(cell, column_widths[col_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _strip_cell_direct_formatting(cell)
            p = cell.paragraphs[0]
            _clear_paragraph_runs(p)
            p.paragraph_format.space_before = Twips(0)
            p.paragraph_format.space_after = Twips(0)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)

    # Enforce minimum row height on ALL rows (0.22" = 317 DXA)
    for row in table.rows:
        row.height = Twips(MIN_ROW_HEIGHT_DXA)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    return table


def _set_gridtable4_tblPr(table, repeat_header=True, last_row_total=False,
                          first_column_emphasis=True):
    """Replace the table's tblPr with the minimal GridTable4 reference block."""
    tbl = table._tbl
    existing = tbl.find(qn("w:tblPr"))
    if existing is not None:
        tbl.remove(existing)

    tblpr = OxmlElement("w:tblPr")
    tbl.insert(0, tblpr)

    ts = OxmlElement("w:tblStyle")
    ts.set(qn("w:val"), "GridTable4")
    tblpr.append(ts)

    tblw = OxmlElement("w:tblW")
    tblw.set(qn("w:w"), "0")
    tblw.set(qn("w:type"), "auto")
    tblpr.append(tblw)

    tbl_look = OxmlElement("w:tblLook")
    tbl_look.set(qn("w:val"), "04A0")
    tbl_look.set(qn("w:firstRow"), "1")
    tbl_look.set(qn("w:lastRow"), "1" if last_row_total else "0")
    tbl_look.set(qn("w:firstColumn"), "1" if first_column_emphasis else "0")
    tbl_look.set(qn("w:lastColumn"), "0")
    tbl_look.set(qn("w:noHBand"), "0")
    tbl_look.set(qn("w:noVBand"), "1")
    tblpr.append(tbl_look)


def _set_tbl_grid(table, column_widths):
    """Set the table's tblGrid."""
    tbl = table._tbl
    existing = tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl.remove(existing)
    tbl_grid = OxmlElement("w:tblGrid")
    tblpr = tbl.find(qn("w:tblPr"))
    if tblpr is not None:
        tblpr.addnext(tbl_grid)
    else:
        tbl.insert(0, tbl_grid)
    for w in column_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        tbl_grid.append(col)


def _clear_paragraph_runs(p):
    """Remove all existing runs from a paragraph, leaving it truly empty.

    python-docx's ``p.text = ""`` clears visible text but can leave a stray
    empty ``<w:r>`` element.  This function removes every ``<w:r>`` child so
    that subsequent ``p.add_run()`` creates the only run in the paragraph.
    """
    for r_el in p._p.findall(qn("w:r")):
        p._p.remove(r_el)


def _strip_cell_direct_formatting(cell):
    """Remove per-cell shading, borders, margins so GridTable4 renders clean."""
    tcpr = cell._tc.find(qn("w:tcPr"))
    if tcpr is None:
        return
    for tag in ("w:shd", "w:tcBorders", "w:tcMar"):
        for el in tcpr.findall(qn(tag)):
            tcpr.remove(el)


def _set_cell_width(cell, width_dxa):
    """Set explicit cell width in DXA."""
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:type"), "dxa")
    tcw.set(qn("w:w"), str(width_dxa))


# ============================================================================
# Non-table direct formatting helpers (callout, cover page only)
# ============================================================================

def _set_cell_shading(cell, color_hex):
    """Cell background fill. Non-table use only (callouts, cover)."""
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)


def _set_cell_borders(cell, color_hex, size=1, style="single"):
    """All four borders on a cell. Non-table use only."""
    tcpr = cell._tc.get_or_add_tcPr()
    tcborders = tcpr.find(qn("w:tcBorders"))
    if tcborders is None:
        tcborders = OxmlElement("w:tcBorders")
        tcpr.append(tcborders)
    for edge in ("top", "left", "bottom", "right"):
        border = tcborders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tcborders.append(border)
        border.set(qn("w:val"), style)
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color_hex)


def _set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    """Cell internal margins in DXA. Non-table use only."""
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        m = tc_mar.find(qn(f"w:{edge}"))
        if m is None:
            m = OxmlElement(f"w:{edge}")
            tc_mar.append(m)
        m.set(qn("w:type"), "dxa")
        m.set(qn("w:w"), str(val))


# ============================================================================
# Cover Page
# ============================================================================

def add_adf_cover_page(doc, title, subtitle=None, client_name=None, date_str=None,
                        prepared_by="ADF Haiti",
                        location="Fond-des-Blancs, Sud, Haiti",
                        full_org_name="Association pour le Développement de Fond-des-Blancs",
                        website="www.adfhaiti.org"):
    """
    Add a cover page matching ADF's standard letterhead design:
    - Prussian Blue band across the top ~35% of the page with white title
    - Subtitle and metadata below the band (left-aligned)
    - Org name and website anchored at the bottom

    Uses a layout table for the blue band (NOT GridTable4; NOT subject to
    data table rules). Call BEFORE any other body content.
    """
    # --- Prussian Blue title band (single-cell table for height control) ---
    # Height: ~3.5" gives roughly 35% of the 9" printable area
    band_table = doc.add_table(rows=1, cols=1)
    band_table.autofit = False
    # Remove all table styling (this is a layout table, not a data table)
    tbl = band_table._tbl
    existing_tblpr = tbl.find(qn("w:tblPr"))
    if existing_tblpr is not None:
        tbl.remove(existing_tblpr)
    tblpr = OxmlElement("w:tblPr")
    tbl.insert(0, tblpr)
    tblw = OxmlElement("w:tblW")
    tblw.set(qn("w:type"), "dxa")
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DEFAULT))
    tblpr.append(tblw)
    # Remove table borders
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tbl_borders.append(b)
    tblpr.append(tbl_borders)

    cell = band_table.rows[0].cells[0]
    _set_cell_width(cell, CONTENT_WIDTH_DEFAULT)
    _set_cell_shading(cell, ADF_COLORS["prussian_blue"])
    _set_cell_margins(cell, top=1800, bottom=600, left=360, right=360)

    # Set row height to ~3.5" (5040 DXA)
    tr = band_table.rows[0]._tr
    trpr = tr.find(qn("w:trPr"))
    if trpr is None:
        trpr = OxmlElement("w:trPr")
        tr.insert(0, trpr)
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), "5040")
    tr_height.set(qn("w:hRule"), "atLeast")
    trpr.append(tr_height)

    # Title text (white, large, bold/SemiBold, left-aligned)
    p_title = cell.paragraphs[0]
    p_title.text = ""
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Twips(0)
    p_title.paragraph_format.space_after = Twips(0)
    r_title = p_title.add_run(title)
    r_title.font.name = ADF_FONT_SEMIBOLD
    r_title.font.size = Pt(28)
    r_title.font.bold = False
    r_title.font.color.rgb = RGBColor.from_string(ADF_COLORS["white"])

    # --- Below the band: subtitle + metadata (left-aligned) ---
    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Twips(200)
    spacer.paragraph_format.space_after = Twips(0)

    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub_p.paragraph_format.space_before = Twips(120)
        sub_p.paragraph_format.space_after = Twips(60)
        sr = sub_p.add_run(subtitle)
        sr.font.name = ADF_FONT_SEMIBOLD
        sr.font.size = Pt(16)
        sr.font.bold = False
        sr.font.color.rgb = RGBColor.from_string(ADF_COLORS["raisin_black"])

    # Descriptive metadata
    meta_lines = []
    if client_name:
        meta_lines.append(f"Prepared for {client_name}")
    meta_lines.append(f"Prepared by {prepared_by}")
    if date_str:
        meta_lines.append(date_str)

    if meta_lines:
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta_p.paragraph_format.space_before = Twips(120)
        meta_p.paragraph_format.space_after = Twips(0)
        meta_text = "\n".join(meta_lines)
        mr = meta_p.add_run(meta_text)
        mr.font.name = ADF_FONT
        mr.font.size = Pt(11)
        mr.font.color.rgb = RGBColor.from_string(ADF_COLORS["davy_grey"])

    # --- Bottom section: org name + website ---
    # Push to bottom with spacer paragraphs
    for _ in range(8):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Twips(0)
        sp.paragraph_format.space_after = Twips(0)

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    org_p.paragraph_format.space_before = Twips(0)
    org_p.paragraph_format.space_after = Twips(40)
    org_r = org_p.add_run(full_org_name)
    org_r.font.name = ADF_FONT
    org_r.font.size = Pt(10)
    org_r.font.color.rgb = RGBColor.from_string(ADF_COLORS["davy_grey"])

    if website:
        web_p = doc.add_paragraph()
        web_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        web_p.paragraph_format.space_before = Twips(0)
        web_p.paragraph_format.space_after = Twips(0)
        web_label = web_p.add_run("Visit our website at ")
        web_label.font.name = ADF_FONT
        web_label.font.size = Pt(10)
        web_label.font.color.rgb = RGBColor.from_string(ADF_COLORS["davy_grey"])
        web_link = web_p.add_run(website)
        web_link.font.name = ADF_FONT
        web_link.font.size = Pt(10)
        web_link.font.color.rgb = RGBColor.from_string(ADF_COLORS["light_sea_green"])

    doc.add_page_break()


def _add_cover_meta(doc, text, size=11):
    """Legacy helper for single-line metadata paragraphs."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Twips(60)
    r = p.add_run(text)
    r.font.name = ADF_FONT
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(ADF_COLORS["davy_grey"])


def _set_paragraph_shading(paragraph, color_hex):
    """Paragraph-level shading for cover page title block."""
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)


# ============================================================================
# Horizontal Rule: FORBIDDEN
# ============================================================================
# ADF explicitly prohibits horizontal rules in documents. The
# insert_horizontal_rule() function has been removed. Do NOT add paragraph
# bottom borders as visual separators in the document body. The only
# permitted paragraph borders are in the header (bottom, Prussian Blue)
# and footer (top, Neptune Cyan).
# ============================================================================


# ============================================================================
# Callout Box (NOT GridTable4)
# ============================================================================

def add_adf_callout(doc, body_text, label=None, total_width=None):
    """
    Callout/info box: single-cell table with Light Sea Green left border,
    Pale Blue Lily fill, Neptune Cyan other borders. Optional bold label.

    NOT a data table; does not use GridTable4.
    """
    if total_width is None:
        total_width = CONTENT_WIDTH_DEFAULT

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False

    tbl = table._tbl
    existing = tbl.find(qn("w:tblPr"))
    if existing is not None:
        tbl.remove(existing)
    tblpr = OxmlElement("w:tblPr")
    tbl.insert(0, tblpr)
    tblw = OxmlElement("w:tblW")
    tblw.set(qn("w:type"), "dxa")
    tblw.set(qn("w:w"), str(total_width))
    tblpr.append(tblw)

    cell = table.rows[0].cells[0]
    _set_cell_width(cell, total_width)
    _set_cell_shading(cell, ADF_COLORS["pale_blue_lily"])
    _set_cell_margins(cell, top=180, bottom=180, left=240, right=180)

    tcpr = cell._tc.get_or_add_tcPr()
    tcborders = OxmlElement("w:tcBorders")
    for edge in ("top", "right", "bottom"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "1")
        b.set(qn("w:color"), ADF_COLORS["neptune_cyan"])
        tcborders.append(b)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:color"), ADF_COLORS["light_sea_green"])
    tcborders.append(left)
    existing_borders = tcpr.find(qn("w:tcBorders"))
    if existing_borders is not None:
        tcpr.remove(existing_borders)
    tcpr.append(tcborders)

    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_before = Twips(0)
    p.paragraph_format.space_after = Twips(0)

    if label:
        lr = p.add_run(label + "  ")
        lr.font.bold = True
        lr.font.color.rgb = RGBColor.from_string(ADF_COLORS["prussian_blue"])

    # Body run inherits font, size, color from Normal style
    br = p.add_run(body_text)

    return table


# ============================================================================
# Signature Block (GridTable4)
# ============================================================================

def add_adf_signature_block(doc, left_label="Siyati Anplwaye", right_label="Siyati Sipevize",
                             row_labels=("Non: ______________________",
                                         "Dat: ______________________",
                                         "Tit: ______________________")):
    """Two-column signature block using GridTable4."""
    table = build_adf_table(
        doc,
        headers=[left_label, right_label],
        rows=[
            ["", ""],
            [row_labels[0], row_labels[0]],
            [row_labels[1], row_labels[1]],
            [row_labels[2], row_labels[2]],
        ],
        column_widths=[5040, 5040],
        first_column_emphasis=False,
    )

    first_data_row = table.rows[1]
    first_data_row.height = Twips(800)
    first_data_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    return table


# ============================================================================
# Paragraph Helper
# ============================================================================

def add_adf_paragraph(doc, text, style=None, bold=False, italic=False,
                      color_hex=None, size_pt=11, alignment=None):
    """Shortcut for adding a properly styled body paragraph.

    By default, runs inherit font, size, and color from the Normal style
    (Aptos 11pt, Raisin Black) set by apply_adf_styles().  Only properties
    that deviate from the style defaults are written to the run, keeping the
    XML clean and letting Word's style engine control alignment and spacing.
    """
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Twips(80)
    run = p.add_run(text)

    # Only set explicit run properties when they differ from Normal style.
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color_hex is not None:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if size_pt != 11:
        run.font.size = Pt(size_pt)
    # font.name intentionally omitted; inherited from Normal style (Aptos).
    return p
